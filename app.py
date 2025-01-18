"""Smart Resume Tailoring App - MVP Implementation."""

import asyncio
import logging
from docx import Document
from dotenv import load_dotenv
from io import BytesIO
import streamlit as st
from typing import Dict

from src.ats_scorer import ATSScorer
from src.job_board import JobBoardClient
from src.job_recommender import JobRecommender
from src.models import ModelManager
from src.salary_analyzer import SalaryAnalyzer

# Configure logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# Load environment variables
load_dotenv()

# Initialize components
model_manager = ModelManager()
ats_scorer = ATSScorer()
salary_analyzer = SalaryAnalyzer()
job_recommender = JobRecommender()


async def process_section(
    section: str,
    section_content: str,
    job_description: str,
    importance: int,
    selected_model: str,
    temperature: float
) -> tuple[str, dict]:
    """Process a single section of the resume."""
    # Generate prompt with importance level
    prompt = model_manager.get_prompt(
        section,
        section_content,
        job_description,
        importance=importance
    )

    # Get completion using selected model
    response = await model_manager.get_completion(
        prompt,
        selected_model,
        temperature
    )

    # Validate output
    is_valid, issues, metrics = model_manager.validate_output(
        section_content,  # Use section content for validation
        response,
        job_description,
        section
    )

    return response, metrics


async def main() -> None:
    """Main application function for the Smart Resume Tailoring App."""
    st.set_page_config(
        page_title="Smart Resume Tailoring App",
        page_icon="📄",
        layout="wide"
    )

    # Load custom CSS
    with open("static/styles.css") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

    # App title with custom styling
    st.markdown(
        """
        <div class="stcard">
            <h1 style='text-align: center; color: #1976D2;'>
                📄 Smart Resume Tailoring App
            </h1>
            <p style='text-align: center; color: #666;'>
                Optimize your resume with AI-powered insights
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    # Sidebar for model selection and configuration
    with st.sidebar:
        st.header("AI Model Configuration")

        # Model selection
        available_models = model_manager.list_available_models()
        model_options = {m["name"]: m["description"] for m in available_models}
        selected_model = st.selectbox(
            "Select AI Model",
            options=list(model_options.keys()),
            help="Choose the AI model for resume tailoring"
        )
        st.caption(model_options[selected_model])

        # Model settings
        temperature = st.slider(
            "Temperature",
            min_value=0.0,
            max_value=1.0,
            value=0.7,
            step=0.1,
            help="Higher values make output more creative, lower more focused"
        )

        # Data collection opt-in
        enable_data_collection = st.checkbox(
            "Allow anonymous data collection",
            value=False,
            help="Help improve the system by sharing anonymized data"
        )

    # Main content area
    # File upload
    uploaded_file = st.file_uploader("Upload your resume (Word)", type=["docx"])

    # Job description input
    job_input_type = st.radio(
        "Job Description Input",
        ["Paste Text", "Enter URL"],
        help="Choose how to input the job description"
    )

    job_description = ""
    if job_input_type == "Paste Text":
        job_description = st.text_area("Paste Job Description")
    else:
        job_url = st.text_input(
            "Enter Job URL",
            help="Enter a Seek.com.au job posting URL"
        )
        if job_url:
            with st.spinner("Fetching job description..."):
                try:
                    client = JobBoardClient()
                    job_details = await client.fetch_job_description(job_url)
                    st.success("Job details fetched successfully!")
                    st.subheader(job_details["title"])
                    st.caption(
                        f"{job_details['company']} - "
                        f"{job_details['location']}"
                    )
                    job_description = job_details["description"]
                    st.text_area(
                        "Fetched Job Description",
                        value=job_description,
                        disabled=True
                    )
                except ValueError as e:
                    st.error(f"Invalid job URL: {str(e)}")
                except Exception as e:
                    st.error(f"Error fetching job description: {str(e)}")

    # Section selection
    st.subheader("Section Selection")
    st.markdown(
        "Choose which sections to tailor and set their importance level. "
        "Higher importance means more focus on matching job requirements."
    )

    # Create columns for section selection
    col1, col2 = st.columns(2)

    # Dictionary to store section selections and their importance
    section_config = {}

    # Left column: Core sections
    with col1:
        st.markdown("### Core Sections")
        if st.checkbox("Summary", value=True):
            section_config["summary"] = st.slider(
                "Summary Importance",
                min_value=1,
                max_value=5,
                value=5,
                help="Higher values mean stronger alignment with job requirements"
            )
        if st.checkbox("Experience", value=True):
            section_config["experience"] = st.slider(
                "Experience Importance",
                min_value=1,
                max_value=5,
                value=4
            )

    # Right column: Additional sections
    with col2:
        st.markdown("### Additional Sections")
        if st.checkbox("Skills", value=True):
            section_config["skills"] = st.slider(
                "Skills Importance",
                min_value=1,
                max_value=5,
                value=4
            )
        if st.checkbox("Education"):
            section_config["education"] = st.slider(
                "Education Importance",
                min_value=1,
                max_value=5,
                value=3
            )

    # Store selected sections
    sections = list(section_config.keys())

    if st.button("Tailor Resume"):
        if uploaded_file and job_description and sections:
            try:
                # Extract text from Word document
                doc = Document(uploaded_file)
                resume_text = "\n".join(p.text for p in doc.paragraphs)

                # Extract sections from resume
                resume_sections: Dict[str, str] = await model_manager._determine_sections(
                    job_description, resume_text
                )
                logger.debug(f"Found resume sections: {list(resume_sections.keys())}")

                # Process selected sections
                tailored_sections = {}
                section_metrics = {}

                # Process each section concurrently
                tasks = []
                valid_sections = []
                for section in sections:
                    section_lower = section.lower()
                    if section_lower in resume_sections:
                        section_content = resume_sections[section_lower]
                        importance = section_config[section]
                        tasks.append(
                            process_section(
                                section,
                                section_content,
                                job_description,
                                importance,
                                selected_model,
                                temperature
                            )
                        )
                        valid_sections.append(section)
                    else:
                        logger.warning(
                            f"Section '{section}' not found in resume. "
                            f"Available sections: {list(resume_sections.keys())}"
                        )
                        st.warning(f"Section '{section}' not found in resume")

                if tasks:  # Only process if we have valid sections
                    # Wait for all sections to be processed
                    results = await asyncio.gather(*tasks)

                    # Store results
                    for section, (content, metrics) in zip(valid_sections, results):
                        tailored_sections[section] = content
                        section_metrics[section] = metrics

                    # Show metrics with custom styling
                    st.markdown(
                        f"""
                        <div class='section-header'>
                            <h4>Metrics for {section}</h4>
                        </div>
                        <div class='metrics-card'>
                            <div style='
                                display: grid;
                                grid-template-columns: repeat(2, 1fr);
                                gap: 1rem;
                            '>
                                <div>
                                    <strong>Length Ratio:</strong>
                                    {metrics['length_ratio']:.2f}
                                </div>
                                <div>
                                    <strong>Keyword Retention:</strong>
                                    {metrics['keyword_retention']:.2f}
                                </div>
                                <div>
                                    <strong>Job Alignment:</strong>
                                    {metrics['job_alignment']:.2f}
                                </div>
                                {
                                    "<div><strong>Metrics Mentioned:</strong> "
                                    + str(metrics['metrics_mentioned']) + "</div>"
                                    if 'metrics_mentioned' in metrics else ""
                                }
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                # Display results with custom styling
                st.markdown(
                    """
                    <div class='section-header'>
                        <h4>Tailored Resume</h4>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                # Create columns for each section
                for section, content in tailored_sections.items():
                    st.markdown(
                        f"""
                        <div class='stcard section-card'>
                            <div class='section-title'>
                                <h3>{section.title()}</h3>
                            </div>
                            <div class='section-content'>{content}</div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

                # Create download options
                st.markdown(
                    """
                    <div class='section-header'>
                        <h4>Download Options</h4>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                col1, col2 = st.columns(2)
                with col1:
                    # Create DOCX in memory
                    doc = Document()
                    for section, content in tailored_sections.items():
                        # Add section header with styling
                        heading = doc.add_heading(section.title(), level=1)
                        from docx.shared import RGBColor
                        heading.style.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)  # Blue color
                        # Add content with formatting
                        paragraph = doc.add_paragraph(content)
                        if paragraph.style:
                            paragraph.style.paragraph_format.space_after = 12

                    # Save to bytes
                    docx_buffer = BytesIO()
                    doc.save(docx_buffer)
                    docx_bytes = docx_buffer.getvalue()

                    st.download_button(
                        label="📥 Download as DOCX",
                        data=docx_bytes,
                        file_name="tailored_resume.docx",
                        mime=(
                            "application/vnd.openxmlformats-"
                            "officedocument.wordprocessingml.document"
                        ),
                        help="Download your tailored resume in Microsoft Word format"
                    )

                with col2:
                    # Create formatted PDF content
                    pdf_sections = []
                    for section, content in tailored_sections.items():
                        pdf_sections.extend([
                            section.upper(),  # Section header in caps
                            "=" * len(section),  # Underline
                            "",  # Blank line
                            content,
                            "\n"  # Extra spacing
                        ])
                    pdf_content = "\n".join(pdf_sections)

                    st.download_button(
                        label="📥 Download as PDF",
                        data=pdf_content.encode(),
                        file_name="tailored_resume.pdf",
                        mime="application/pdf",
                        help="Download your tailored resume in PDF format"
                    )

                # Collect training data if opted in
                if enable_data_collection and tailored_sections:
                    try:
                        # Collect training data concurrently
                        training_tasks: list[asyncio.Task[None]] = []
                        for section, content in tailored_sections.items():
                            if section in section_metrics:  # Only collect if we have metrics
                                # Create task for each training data collection
                                coro = model_manager.collect_training_data(
                                    resume_text,
                                    job_description,
                                    content,
                                    section,
                                    section_metrics[section]
                                )
                                task = asyncio.create_task(coro)
                                training_tasks.append(task)
                        if training_tasks:
                            await asyncio.gather(*training_tasks)
                    except Exception as e:
                        st.warning(f"Failed to collect training data: {str(e)}")

            except Exception as e:
                st.error(f"Error processing resume: {str(e)}")
        else:
            st.warning(
                "Please upload a resume, enter a job description, "
                "and select sections to tailor."
            )


if __name__ == "__main__":
    asyncio.run(main())
