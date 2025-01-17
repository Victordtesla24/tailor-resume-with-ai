"""Tests for document formatting preservation."""

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from src.exceptions import FormatError
from src.formatting import FormatPreserver, StyleMap


@pytest.fixture
def sample_doc_path(tmp_path: Path) -> str:
    """Create a sample document for testing."""
    doc = Document()

    # Add some styled content
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]
    run = para.add_run("Test content with ")
    run.font.size = Pt(12)
    run = para.add_run("different ")
    run.font.size = Pt(14)
    run.bold = True
    run = para.add_run("styles")
    run.font.size = Pt(12)
    run.italic = True

    # Save document
    doc_path = tmp_path / "test.docx"
    doc.save(str(doc_path))
    return str(doc_path)


def test_style_map_initialization():
    """Test StyleMap initialization."""
    style_map = StyleMap()
    assert isinstance(style_map.paragraph_styles, dict)
    assert isinstance(style_map.run_styles, dict)


def test_style_map_operations():
    """Test StyleMap add and get operations."""
    style_map = StyleMap()

    # Test paragraph style operations
    para_style = {"alignment": 1, "style_name": "Normal"}
    style_map.add_paragraph_style(0, para_style)
    assert style_map.get_paragraph_style(0) == para_style
    assert style_map.get_paragraph_style(1) is None

    # Test run style operations
    run_style = {"bold": True, "font_size": 12}
    style_map.add_run_style(0, run_style)
    assert style_map.get_run_styles(0) == [run_style]
    assert style_map.get_run_styles(1) == []


def test_format_preserver_initialization(sample_doc_path: str):
    """Test FormatPreserver initialization."""
    preserver = FormatPreserver(sample_doc_path)
    assert preserver.doc is not None
    assert isinstance(preserver.style_map, StyleMap)


def test_format_preserver_style_extraction(sample_doc_path: str):
    """Test style extraction from document."""
    preserver = FormatPreserver(sample_doc_path)
    style_map = preserver.style_map

    # Check paragraph styles
    para_style = style_map.get_paragraph_style(0)
    assert para_style is not None
    assert para_style["style_name"] == "Normal"

    # Check run styles
    run_styles = style_map.get_run_styles(0)
    assert len(run_styles) == 3
    assert run_styles[1]["bold"] is True
    assert run_styles[2]["italic"] is True


def test_format_preserver_style_application(sample_doc_path: str):
    """Test applying styles to updated content."""
    preserver = FormatPreserver(sample_doc_path)

    # Update content while preserving styles
    updated_sections = {0: "Updated content with preserved styling"}
    new_doc = preserver.apply_styles(updated_sections)

    # Verify the new document
    assert len(new_doc.paragraphs) > 0
    first_para = new_doc.paragraphs[0]
    assert first_para.text == "Updated content with preserved styling"
    if first_para.style:
        assert first_para.style.name == "Normal"


def test_format_preserver_error_handling():
    """Test error handling for invalid documents."""
    with pytest.raises(FormatError):
        FormatPreserver("nonexistent.docx")


def test_style_preservation_with_empty_update(sample_doc_path: str):
    """Test style preservation with no content updates."""
    preserver = FormatPreserver(sample_doc_path)
    new_doc = preserver.apply_styles({})

    # Verify original content and styles are preserved
    assert len(new_doc.paragraphs) == len(preserver.doc.paragraphs)
    for orig_para, new_para in zip(preserver.doc.paragraphs, new_doc.paragraphs):
        assert orig_para.text == new_para.text
        if orig_para.style and new_para.style:
            assert orig_para.style.name == new_para.style.name
