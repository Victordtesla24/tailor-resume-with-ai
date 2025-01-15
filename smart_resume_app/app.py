"""Smart Resume Tailoring Application."""
import logging
from pathlib import Path
from typing import Optional, Dict, Any
from docx import Document
from dotenv import load_dotenv
import streamlit as st

from src.components import (
    ComparisonView,
    ProgressIndicator,
    SectionSelector,
    setup_custom_styling
)
from src.config import Config, load_config
from src.data_collection import DataCollector
from src.formatting import FormatPreserver
from src.job_board import JobBoardScraper
from src.exceptions import JobBoardError
from src.keyword_matcher import (
    KeywordMatcher,
    SectionDetector,
    SkillExtractor
)
from src.models import AIModelManager

# Load environment variables
load_dotenv()

# Setup logging
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger('resume_tailor')

# Initialize components
config = Config()
config.load_env_vars()

# Load app settings
app_settings = load_config()

# Initialize managers and components
ai_manager = AIModelManager(config)
job_board = JobBoardScraper()
comparison_view = ComparisonView()
progress = ProgressIndicator()
section_selector = SectionSelector()
skill_extractor = SkillExtractor()
data_collector = DataCollector()


def init_session_state():
    """Initialize session state variables."""
    if 'uploaded_file' not in st.session_state:
        st.session_state.uploaded_file = None
    if 'job_description' not in st.session_state:
        st.session_state.job_description = None
    if 'selected_model' not in st.session_state:
        st.session_state.selected_model = 'gpt-4'
    if 'tailored_content' not in st.session_state:
        st.session_state.tailored_content = None


def save_uploaded_file(uploaded_file) -> Optional[str]:
    """Save uploaded file and return path."""
    try:
        save_path = Path("uploads") / uploaded_file.name
        save_path.parent.mkdir(exist_ok=True)

        with open(save_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        return str(save_path)
    except Exception as e:
        logger.error(f"Error saving file: {str(e)}")
        return None


def process_resume(
    file_path: str,
    job_description: str,
    selected_sections: list,
    model_name: str
) -> Dict[str, Any]:
    """Process resume with selected AI model."""
    try:
        # Initialize progress
        progress.start()

        # Extract resume content
        doc = Document(file_path)
        resume_text = "\n".join(p.text for p in doc.paragraphs)
        progress.next_step()

        # Extract job requirements
        keyword_matcher = KeywordMatcher()
        skill_scores = keyword_matcher.match_skills(
            resume_text,
            job_description
        )
        progress.next_step()

        # Switch to selected model
        ai_manager.switch_model(model_name)

        # Process with AI
        tailored_text = ai_manager.process_resume(
            resume_text,
            job_description,
            selected_sections
        )
        progress.next_step()

        # Preserve formatting
        format_preserver = FormatPreserver(file_path)
        output_path = format_preserver.apply_styles({
            0: tailored_text
        })
        progress.next_step()

        # Save anonymized training data
        if st.session_state.get('share_data', False):
            data_collector.save_training_data(
                resume_text,
                job_description,
                tailored_text,
                {"model": model_name}
            )

        return {
            'success': True,
            'tailored_text': tailored_text,
            'output_path': output_path,
            'skill_scores': skill_scores
        }

    except Exception as e:
        logger.error(f"Error processing resume: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }


def main():
    """Main application."""
    st.set_page_config(page_title="Smart Resume Tailoring", layout="wide")

    setup_custom_styling()
    init_session_state()

    st.title("Smart Resume Tailoring")

    # Sidebar
    with st.sidebar:
        st.header("Settings")

        # Model selection
        st.session_state.selected_model = st.selectbox(
            "Select AI Model",
            options=ai_manager.list_available_models()
        )

        # Data sharing opt-in
        share_text = "Share anonymized data for improvement"
        st.session_state.share_data = st.checkbox(share_text, value=False)

    # Main content
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Upload Resume")
        uploaded_file = st.file_uploader("Choose a file", type=['docx'])

        if uploaded_file:
            st.session_state.uploaded_file = save_uploaded_file(uploaded_file)

    with col2:
        st.subheader("Job Description")
        job_url = st.text_input("Enter job URL (optional)")

        if job_url:
            with st.spinner("Fetching job description..."):
                try:
                    job_description = job_board.fetch_job_description(job_url)
                    if job_description:
                        st.session_state.job_description = job_description
                        st.success("Job description fetched successfully")
                except JobBoardError as e:
                    st.error(
                        "Could not fetch job description. "
                        "Please paste it manually below.\n\n"
                        f"Error: {str(e)}"
                    )
                    # Show supported job boards
                    supported = job_board.get_supported_sites()
                    st.info(
                        "Supported job boards: " + ", ".join(supported)
                    )

        text_area_label = "Or paste job description"
        text_area_value = st.session_state.job_description or ""
        st.session_state.job_description = st.text_area(
            text_area_label, value=text_area_value, height=200
        )

    # Section selection
    if st.session_state.uploaded_file:
        section_detector = SectionDetector()
        sections = section_detector.identify_sections(
            Document(st.session_state.uploaded_file)
        )
        selected_sections = section_selector.display_selector(sections)

    # Process button
    if (
        st.button("Tailor Resume")
        and st.session_state.uploaded_file
        and st.session_state.job_description
    ):

        with st.spinner("Processing..."):
            result = process_resume(
                st.session_state.uploaded_file,
                st.session_state.job_description,
                selected_sections,
                st.session_state.selected_model
            )

            if result['success']:
                st.session_state.tailored_content = result['tailored_text']

                # Display comparison
                comparison_view.display_comparison(
                    Document(st.session_state.uploaded_file),
                    result['tailored_text']
                )

                # Download button
                with open(result['output_path'], 'rb') as f:
                    st.download_button(
                        "Download Tailored Resume",
                        f,
                        file_name="Tailored_Resume.docx"
                    )

                # Show skill matches
                st.subheader("Skill Match Analysis")
                for skill, score in result['skill_scores'].items():
                    st.progress(
                        score,
                        text=f"{skill}: {score:.0%}"
                    )
            else:
                st.error(f"Error processing resume: {result['error']}")


if __name__ == "__main__":
    main()
