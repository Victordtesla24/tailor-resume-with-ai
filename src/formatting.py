"""Resume formatting preservation functionality."""

import logging
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt

from src.exceptions import FormatError

# Setup logging
logger = logging.getLogger(__name__)


class StyleMap:
    """Map of document styles and formatting."""

    def __init__(self) -> None:
        """Initialize style map."""
        self.paragraph_styles: Dict[int, Dict[str, Any]] = {}
        self.run_styles: Dict[int, List[Dict[str, Any]]] = {}

    def add_paragraph_style(self, index: int, style: Dict[str, Any]) -> None:
        """Add paragraph style to map."""
        self.paragraph_styles[index] = style

    def add_run_style(self, para_index: int, style: Dict[str, Any]) -> None:
        """Add run style to map."""
        if para_index not in self.run_styles:
            self.run_styles[para_index] = []
        self.run_styles[para_index].append(style)

    def get_paragraph_style(self, index: int) -> Optional[Dict[str, Any]]:
        """Get paragraph style by index."""
        return self.paragraph_styles.get(index)

    def get_run_styles(self, para_index: int) -> List[Dict[str, Any]]:
        """Get run styles for paragraph."""
        return self.run_styles.get(para_index, [])


class FormatPreserver:
    """Preserve document formatting during tailoring."""

    def __init__(self, doc_path: str):
        """Initialize with document path."""
        try:
            self.doc = Document(doc_path)
            self.style_map = self._extract_styles()
        except Exception as e:
            logger.error(f"Error loading document: {str(e)}")
            raise FormatError(f"Failed to load document: {str(e)}")

    def _extract_styles(self) -> StyleMap:
        """Extract styles from document."""
        style_map = StyleMap()

        for idx, para in enumerate(self.doc.paragraphs):
            # Extract paragraph style
            style_name = getattr(para.style, "name", None)
            para_style = {
                "alignment": para.alignment,
                "style_name": style_name,
                "space_before": (
                    para.paragraph_format.space_before
                    if para.paragraph_format.space_before
                    else None
                ),
                "space_after": (
                    para.paragraph_format.space_after if para.paragraph_format.space_after else None
                ),
                "line_spacing": (
                    para.paragraph_format.line_spacing
                    if para.paragraph_format.line_spacing
                    else None
                ),
            }
            style_map.add_paragraph_style(idx, para_style)

            # Extract run styles
            for run in para.runs:
                run_style = {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name,
                    "font_size": (run.font.size.pt if run.font.size else None),
                    "color": (run.font.color.rgb if run.font.color else None),
                }
                style_map.add_run_style(idx, run_style)

        return style_map

    def apply_styles(
        self, updated_sections: Dict[int, str]
    ) -> Any:  # Return type is docx.document.Document
        """Apply preserved styles to updated content."""
        try:
            # Create new document
            new_doc = Document()

            # Copy styles from original
            for para_idx, para in enumerate(self.doc.paragraphs):
                new_para = new_doc.add_paragraph()

                # Apply paragraph style
                if style := self.style_map.get_paragraph_style(para_idx):
                    if style["alignment"]:
                        new_para.alignment = style["alignment"]
                    style_name = style.get("style_name")
                    if style_name:
                        new_para.style = style_name
                    if style["space_before"]:
                        new_para.paragraph_format.space_before = style["space_before"]
                    if style["space_after"]:
                        new_para.paragraph_format.space_after = style["space_after"]
                    if style["line_spacing"]:
                        new_para.paragraph_format.line_spacing = style["line_spacing"]

                # Get updated or original text
                text = updated_sections.get(para_idx, para.text)

                # Apply run styles
                run_styles = self.style_map.get_run_styles(para_idx)
                if run_styles:
                    # Try to match original style segments
                    segments = self._split_by_style_segments(text, run_styles)
                    for segment, style in zip(segments, run_styles):
                        run = new_para.add_run(segment)
                        self._apply_run_style(run, style)
                else:
                    # No style segments, add as single run
                    run = new_para.add_run(text)

            return new_doc

        except Exception as e:
            logger.error(f"Error applying styles: {str(e)}")
            raise FormatError(f"Failed to apply styles: {str(e)}")

    def _split_by_style_segments(self, text: str, styles: List[Dict[str, Any]]) -> List[str]:
        """Split text to match original style segments."""
        if not styles:
            return [text]

        # Try to match original segments
        total_len = sum(len(style["text"]) for style in styles)
        if total_len == 0:
            return [text]

        # Split proportionally
        segments = []
        remaining = text
        for i, style in enumerate(styles[:-1]):
            ratio = len(style["text"]) / total_len
            split_point = int(len(text) * ratio)
            segments.append(remaining[:split_point])
            remaining = remaining[split_point:]
        segments.append(remaining)

        return segments

    def _apply_run_style(self, run: Any, style: Dict[str, Any]) -> None:
        """Apply style to run."""
        if style["bold"] is not None:
            run.bold = style["bold"]
        if style["italic"] is not None:
            run.italic = style["italic"]
        if style["underline"] is not None:
            run.underline = style["underline"]
        if style["font_name"]:
            run.font.name = style["font_name"]
        if style["font_size"]:
            run.font.size = Pt(style["font_size"])
        if style["color"]:
            run.font.color.rgb = style["color"]
